from flask import Flask, request, send_file, jsonify
import os
import sys
import tempfile
import shutil
import fitz  # PyMuPDF
from docx import Document
import time
from werkzeug.utils import secure_filename

# Constants
HOST = "0.0.0.0"  # Listen on all interfaces
PORT = 8081  # Use a different port to avoid conflicts
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB limit

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = Flask(__name__)

def pdf_to_word(pdf_path, docx_path):
    """Convert PDF to DOCX using PyMuPDF and python-docx"""
    doc = Document()
    pdf = fitz.open(pdf_path)
    
    # Add a cover page with info
    doc.add_heading('PDF to Word Conversion', 0)
    doc.add_paragraph(f'Original filename: {os.path.basename(pdf_path)}')
    doc.add_paragraph(f'Total pages: {len(pdf)}')
    doc.add_page_break()
    
    # Process each page
    for page_num in range(len(pdf)):
        # Add a page header
        doc.add_heading(f'Page {page_num + 1}', level=1)
        
        # Get the page
        page = pdf.load_page(page_num)
        
        # Extract text blocks to preserve some structure
        blocks = page.get_text("blocks")
        for block in blocks:
            # The last item in the block tuple is the text
            text = block[4]
            # Add the text as a paragraph
            doc.add_paragraph(text)
        
        # Add a page break after each page except the last one
        if page_num < len(pdf) - 1:
            doc.add_page_break()
    
    # Save the document
    doc.save(docx_path)
    return True

@app.route('/status', methods=['GET', 'OPTIONS'])
def status():
    if request.method == 'OPTIONS':
        return '', 200
    return jsonify({
        'status': 'running',
        'service': 'PDF to DOCX Conversion Server',
        'version': '1.0.0'
    })

@app.route('/convert-pdf-to-docx', methods=['POST', 'OPTIONS'])
def convert():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Check if file is in the request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        # Get the file
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check if it's a PDF
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'File must be a PDF'}), 400
        
        # Create temporary paths
        timestamp = int(time.time())
        secure_name = secure_filename(file.filename)
        temp_pdf = os.path.join(UPLOAD_DIR, f"{timestamp}_{secure_name}")
        temp_docx = os.path.join(UPLOAD_DIR, f"{timestamp}_{secure_name.replace('.pdf', '.docx')}")
        
        # Save the uploaded file
        file.save(temp_pdf)
        
        # Check file size
        if os.path.getsize(temp_pdf) > MAX_FILE_SIZE:
            os.remove(temp_pdf)
            return jsonify({'error': f'File too large (max {MAX_FILE_SIZE/1024/1024}MB)'}), 413
        
        # Convert PDF to DOCX
        try:
            print(f"Converting {temp_pdf} to {temp_docx}")
            pdf_to_word(temp_pdf, temp_docx)
            
            # Check if conversion was successful
            if not os.path.exists(temp_docx) or os.path.getsize(temp_docx) == 0:
                raise Exception("Conversion failed to produce a valid DOCX file")
            
            # Send the file
            response = send_file(
                temp_docx,
                as_attachment=True,
                download_name=os.path.basename(temp_docx),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            # Set headers for CORS
            response.headers.add('Access-Control-Allow-Origin', '*')
            
            # Clean up files after sending response
            @response.call_on_close
            def cleanup():
                try:
                    if os.path.exists(temp_pdf):
                        os.remove(temp_pdf)
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                except Exception as e:
                    print(f"Cleanup error: {str(e)}")
                    
            return response
            
        except Exception as e:
            # Clean up files on error
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            print(f"Conversion error: {str(e)}")
            return jsonify({'error': f'Conversion error: {str(e)}'}), 500
        
    except Exception as e:
        print(f"Server error: {str(e)}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/convert-pdf-to-docx', methods=['OPTIONS'])
def options():
    """Handle CORS preflight requests"""
    response = app.make_default_options_response()
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    return response

if __name__ == '__main__':
    try:
        print(f"PDF Conversion Server running at http://localhost:{PORT}")
        print("Use the /convert-pdf-to-docx endpoint for PDF to DOCX conversion")
        print("Press Ctrl+C to stop the server")
        app.run(host=HOST, port=PORT, debug=False)
    except OSError as e:
        if getattr(e, 'errno', None) == 10048:  # Port already in use
            print(f"Port {PORT} is already in use. Choose a different port or stop the existing server.")
        else:
            print(f"Error starting server: {e}")
    except KeyboardInterrupt:
        print("\nServer stopped by user")
        sys.exit(0) 